import os
from docx import Document
from lightrag import LightRAG, QueryParam
from lightrag.llm import ollama_model_complete, ollama_embedding
from lightrag.utils import EmbeddingFunc
import pandas as pd
import textract

WORKING_DIR = "./out"

if not os.path.exists(WORKING_DIR):
    os.mkdir(WORKING_DIR)

rag = LightRAG(
    working_dir=WORKING_DIR,
    llm_model_func=ollama_model_complete,
    llm_model_name="qwen2m",
    
    embedding_func=EmbeddingFunc(
        embedding_dim=768,
        max_token_size=8192,
        func=lambda texts: ollama_embedding(texts, embed_model="nomic-embed-text:latest"),
    ),
)

'''

# Open the .docx file
doc = Document("C:\\Users\\jmart\\Downloads\\ChatCat-main\\model\\SFWE-Graduate-Student-Handbook.docx")
doc_text = "\n".join([para.text for para in doc.paragraphs])
rag.insert(doc_text)

doc = Document("C:\\Users\\jmart\\Downloads\\Main SFWE Degree Planning Guide.docx")
doc_text = "\n".join([para.text for para in doc.paragraphs])
rag.insert(doc_text)

with open("C:\\Users\\jmart\\Downloads\\SIE-Undergraduate-Handbook-23-24.txt", encoding="utf8") as f:
    rag.insert(f.read())

with open("C:\\Users\\jmart\\Downloads\\SoftwareEngineering24-25-1.txt", encoding="utf8") as f:
    rag.insert(f.read())

with open("C:\\Users\\jmart\\Downloads\\program-SFEPHD.txt", encoding="utf8") as f:
    rag.insert(f.read())

df = pd.read_excel("C:\\Users\\jmart\\Downloads\\SFWE Tech Electives 2024 100824.xlsx")
rag.insert(df)

df = pd.read_excel("C:\\Users\\jmart\\Downloads\\SFWE Grad Tech Electives.xlsx")
rag.insert(df)

with open("C:\\Users\\jmart\\Downloads\\Understanding Your Financial Aid Offer _ Office of Scholarships & Financial Aid.txt", encoding="utf8") as f:
    rag.insert(f.read())





with open("C:\\Users\\jmart\\Downloads\\Welcome _ Office of Scholarships & Financial Aid.txt", encoding="utf8") as f:
    rag.insert(f.read())

with open("C:\\Users\\jmart\\Downloads\\financialaidwebsite.txt", encoding="utf8") as f:
    rag.insert(f.read())


with open("C:\\Users\\jmart\\Downloads\\Research _ Software Engineering.txt", encoding="utf8") as f:
    rag.insert(f.read())

with open("C:\\Users\\jmart\\Downloads\\Home _ Research, Innovation, and Impact.txt", encoding="utf8") as f:
    rag.insert(f.read())

with open("C:\\Users\\jmart\\Downloads\\About _ Research, Innovation, and Impact.txt", encoding="utf8") as f:
    rag.insert(f.read())
'''
with open("C:\\Users\\jmart\\Downloads\\UA-Software-Engineering-Degree-_-University-of-Arizona-College-of-Engineering.txt", encoding="utf8") as f:
    rag.insert(f.read())



print(
    rag.query("What are the research opportunities for software engineering at the university of arizona?", param=QueryParam(mode="hybrid"))
)
'''
print(
    rag.query("what are the admission requirements for the software engineering doctoral program?", param=QueryParam(mode="hybrid"))
)
# Perform naive search

print("Asking question 1-1...")
print(
    rag.query("If I have a GPA of 2.85 do I meet Satisfactory Academic Progress for the undergraduate program?", param=QueryParam(mode="hybrid"))
)

print("Asking question 1-2...")
print(
    rag.query("If I have a GPA of 2.85 do I meet Satisfactory Academic Progress for the graduate program?", param=QueryParam(mode="hybrid"))
)

print("Asking question 2...")
print(
    rag.query("What website should I visit to apply for Graduate Admissions?", param=QueryParam(mode="hybrid"))
)

print("Asking question 3...")
print(
    rag.query("What are the Graduate Admission Requirements?", param=QueryParam(mode="hybrid"))
)

print("Asking question 4...")
print(
    rag.query("What are some Electives for Software Engineering?", param=QueryParam(mode="hybrid"))
)

'''